import streamlit as st
import openai
import PyPDF2
import re
from io import BytesIO
from docx import Document
from googleapiclient.discovery import build
import google.generativeai as genai  # Import Gemini API
from googleapiclient.errors import HttpError
import os

# ✅ Prevent Google APIs from using default credentials (forces API Key usage)
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = ""

# ✅ Google Gemini API Key
try:
    GEMINI_API_KEY = "AIzaSyC-huBpwrdLt2y3Cabi3pVxa6gmCX0y2fU"
    genai.configure(api_key=GEMINI_API_KEY)
except Exception as e:
    st.error(f"Gemini API Error: {e}")

# ✅ YouTube API Key
try:
    YOUTUBE_API_KEY = "AIzaSyAz4ecTJB8g7GKTm5zQkxNc09sm9qeBXGo"
    youtube = build("youtube", "v3", developerKey=YOUTUBE_API_KEY)  # Ensure API Key is used
except HttpError as e:
    st.error(f"YouTube API Error: {e}")
except Exception as e:
    st.error(f"Unexpected error in YouTube API: {e}")

# ✅ OpenAI (SambaNova AI) API Key
try:
    OPENAI_API_KEY = "e146e13e-46f6-4d38-8b67-a120c374803b"
    client = openai.OpenAI(
        api_key=OPENAI_API_KEY,
        base_url="https://api.sambanova.ai/v1",
    )
except Exception as e:
    st.error(f"OpenAI API Error: {e}")


def get_chatbot_response(messages, client):
    try:
        response = client.chat.completions.create(
            model="Meta-Llama-3.3-70B-Instruct",
            messages=messages,
            temperature=0.1,
            top_p=0.1,
            max_tokens=4096
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error: {str(e)}"


def extract_text_from_pdf(uploaded_file):
    pdf_reader = PyPDF2.PdfReader(uploaded_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text


def extract_titles_and_contents(text):
    sections = re.split(r'(?m)^(?=[A-Z][^\n]*$)', text)
    structured_content = {}
    for section in sections:
        lines = section.strip().split("\n", 1)
        if len(lines) == 2:
            title, content = lines
            structured_content[title.strip()] = content.strip()
    return structured_content


def generate_word_doc(explanations):
    doc = Document()
    doc.add_heading("Detailed Explanations", level=1)
    for title, content in explanations.items():
        cleaned_content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # Remove **bold** formatting
        cleaned_content = re.sub(r'\*(.*?)\*', r'\1', cleaned_content)  # Remove *italic* formatting
        cleaned_content = re.sub(r'__(.*?)__', r'\1', cleaned_content)  # Remove __underline__ formatting
        cleaned_content = re.sub(r'_(.*?)_', r'\1', cleaned_content)  # Remove _italic_ formatting
        doc.add_heading(title, level=2)
        doc.add_paragraph(cleaned_content)
    word_output = BytesIO()
    doc.save(word_output)
    word_output.seek(0)
    return word_output


def search_videos(content):
    search_query = f"What is {content}"
    try:
        search_response = youtube.search().list(
            q=search_query,
            part="snippet",
            type="video",
            maxResults=5
        ).execute()
        videos = []
        for item in search_response["items"]:
            videos.append({
                "title": item["snippet"]["title"],
                "description": item["snippet"]["description"],
                "thumbnail": item["snippet"]["thumbnails"]["high"]["url"],
                "video_id": item["id"]["videoId"]
            })
        return videos
    except Exception as e:
        st.error(f"Error fetching YouTube videos: {e}")
        return []


def get_study_material_links(topic):
    """Fetch relevant study material links (open access) using SambaNova API."""
    # Create a prompt that specifically asks for free academic resources
    prompt = f"Provide open-access academic resources (preferably PDFs or free textbooks) for the topic: {topic}. Include links from sites like arXiv, Google Scholar, Project Gutenberg, or ResearchGate."

    try:
        # Create the request for the SambaNova model
        response = client.chat.completions.create(
            model="Meta-Llama-3.3-70B-Instruct",  # Specify the model you wish to use
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            top_p=0.1,
            max_tokens=4096
        )
        
        # Extract the response text from SambaNova API
        study_links = response.choices[0].message.content.strip().split("\n")
        
        # Filter only valid HTTP links (HTTP, HTTPS)
        valid_links = [link.strip() for link in study_links if "http" in link]

        # Further refine to include only links to open-access academic sites
        open_access_links = []
        for link in valid_links:
            if any(domain in link for domain in ["arxiv.org", "scholar.google.com", "gutenberg.org", "researchgate.net"]):
                open_access_links.append(link)

        # Fallback if no open-access links are found
        if not open_access_links:
            open_access_links = [
                f"🔗 [Free Resource on Google Scholar](https://scholar.google.com/scholar?q={topic})",
                f"🔗 [Free PDFs on arXiv](https://arxiv.org/search/?query={topic}&searchtype=all)",
                f"🔗 [Books on Project Gutenberg](https://www.gutenberg.org/ebooks/search/?query={topic})",
            ]

        return open_access_links
    
    except Exception as e:
        return [f"Error fetching links: {str(e)}"]



def main():
    st.set_page_config(page_title="NVidya", layout="wide")
    st.markdown("<h1 style='text-align: center; color: yellow;'>NVidya - Your Intelligent Learning Partner</h1>",
                unsafe_allow_html=True)

    if "messages" not in st.session_state:
        st.session_state.messages = [
            {"role": "system",
             "content": "You are a detailed and knowledgeable assistant. Provide thorough explanations."}
        ]

    st.sidebar.header("Input Fields")
    user_title = st.sidebar.text_input("Enter Title")
    uploaded_file = st.sidebar.file_uploader("Upload PDF", type=["pdf"])

    explanations = {}

    if uploaded_file is not None:
        extracted_text = extract_text_from_pdf(uploaded_file)
        structured_content = extract_titles_and_contents(extracted_text)

        for title, content in structured_content.items():
            user_prompt = f"Provide an in-depth study material explanation for: {title}. Ensure it follows a structured format and is detailed. Also do day wise split for each topic, like in day 1 this topic can be covered. Add codings and formulas in the generated content if required to the topic."
            st.session_state.messages.append({"role": "user", "content": user_prompt})
            bot_response = get_chatbot_response(st.session_state.messages, client)
            explanations[title] = bot_response
            st.session_state.messages.append({"role": "assistant", "content": bot_response})

    if user_title:
        user_prompt = f"Title: {user_title}\nProvide a detailed and thorough explanation."
        st.session_state.messages.append({"role": "user", "content": user_prompt})
        bot_response = get_chatbot_response(st.session_state.messages, client)
        explanations[user_title] = bot_response
        st.session_state.messages.append({"role": "assistant", "content": bot_response})

    if explanations:
        # First, display the explanations
        for title, explanation in explanations.items():
            st.write(explanation)

        # Then, show the download button
        word_file = generate_word_doc(explanations)
        st.download_button(
            label="Download Explanations as Word Document",
            data=word_file,
            file_name="explanations.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Fetch study material links
        st.subheader("📚 Study Material Links")
        study_links = get_study_material_links(title)
        for link in study_links:
            st.markdown(f"- 🔗 [Study Material]({link})" if "http" in link else f"- ❌ {link}")

        # Then display the recommended YouTube videos
        st.subheader("Recommended YouTube Videos 📺")
        with st.spinner("Fetching relevant YouTube videos..."):
            videos = search_videos(user_title if user_title else "AI learning")
            if videos:
                for video in videos:
                    with st.expander(video["title"]):
                        st.image(video["thumbnail"], width=200)
                        st.markdown(f"**[{video['title']}](https://www.youtube.com/watch?v={video['video_id']})**")
                        st.write(video["description"])
            else:
                st.error("No related videos found.")

    else:
        st.write("No topics extracted or explained yet.")


if __name__ == "__main__":
    main()
